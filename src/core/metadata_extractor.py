"""
Auto Metadata Extraction Module (Feature 40)

Extracts metadata (author, date, title, etc.) from uploaded documents.
Supports PDF, DOCX, and text files via regex patterns and library-specific APIs.

Usage:
    from src.core.metadata_extractor import extract_metadata

    meta = extract_metadata("path/to/document.pdf")
    print(meta.title, meta.author, meta.created_date)
"""

import datetime
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from src.core import setup_logger

logger = setup_logger("metadata")


# ============================================================
# Data Model
# ============================================================


@dataclass
class DocumentMetadata:
    """Extracted metadata from a document.

    Attributes:
        filename: Original filename
        file_extension: File extension (lowercase, no dot)
        file_size_bytes: File size
        title: Document title (extracted or inferred)
        author: Document author (if available)
        created_date: Creation date string (if available)
        modified_date: Last modified date string
        page_count: Number of pages (if available, PDF/DOCX)
        word_count: Word count (if available)
        language: Detected language (if available)
        keywords: List of extracted keywords/tags
        mime_type: MIME type
        raw: Raw metadata dict from the library
    """

    filename: str = ""
    file_extension: str = ""
    file_size_bytes: int = 0
    title: str = ""
    author: str = ""
    created_date: str = ""
    modified_date: str = ""
    page_count: int = 0
    word_count: int = 0
    language: str = ""
    keywords: list[str] = field(default_factory=list)
    mime_type: str = ""
    raw: dict = field(default_factory=dict)


# ============================================================
# Extraction Functions
# ============================================================

_TITLE_PATTERNS = [
    re.compile(r"^#\s+(.+)$", re.MULTILINE),  # Markdown H1
    re.compile(r"^(.+)\n[=]+\s*$", re.MULTILINE),  # Underlined title
    re.compile(r"<title>(.+?)</title>", re.IGNORECASE | re.DOTALL),  # HTML title
]

_AUTHOR_PATTERNS = [
    re.compile(r"^(?:Author|Tác giả|By|Written by)[:\s]+(.+)$", re.MULTILINE | re.IGNORECASE),
    re.compile(r"@author\s+(.+)", re.IGNORECASE),
]

_DATE_PATTERNS = [
    re.compile(r"^(?:Date|Created|Ngày)[:\s]+(.+)$", re.MULTILINE | re.IGNORECASE),
    re.compile(r"(\d{4}-\d{2}-\d{2})"),
    re.compile(r"(\d{2}/\d{2}/\d{4})"),
]

_KEYWORD_PATTERNS = [
    re.compile(r"^(?:Tags|Keywords|Từ khóa)[:\s]+(.+)$", re.MULTILINE | re.IGNORECASE),
]


def extract_metadata(filepath: str) -> DocumentMetadata:
    """Extract metadata from a document file.

    Supports PDF, DOCX, TXT, MD, CSV, JSON, HTML files.
    For unsupported formats, returns basic file metadata.

    Args:
        filepath: Path to the document

    Returns:
        DocumentMetadata with extracted fields
    """
    path = Path(filepath)
    meta = DocumentMetadata(
        filename=path.name,
        file_extension=path.suffix.lower().lstrip("."),
        file_size_bytes=path.stat().st_size if path.exists() else 0,
        modified_date=_format_timestamp(path.stat().st_mtime) if path.exists() else "",
        mime_type=_guess_mime_type(path.suffix),
    )

    if not path.exists():
        return meta

    ext = meta.file_extension

    try:
        if ext == "pdf":
            _extract_pdf_metadata(path, meta)
        elif ext == "docx":
            _extract_docx_metadata(path, meta)
        elif ext in ("txt", "md", "csv", "json", "html", "htm", "xml", "py", "js", "ts"):
            _extract_text_metadata(path, meta)
    except Exception as e:
        logger.debug(f"Metadata extraction failed for {filepath}: {e}")

    return meta


def _extract_pdf_metadata(path: Path, meta: DocumentMetadata) -> None:
    """Extract metadata from PDF using PyPDF2 or pikepdf."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(str(path))
        info = reader.metadata or {}
        meta.author = info.get("/Author", "") or ""
        meta.title = info.get("/Title", "") or ""
        meta.created_date = str(info.get("/CreationDate", ""))[:10] if info.get("/CreationDate") else ""
        meta.page_count = len(reader.pages)
        meta.raw = {k: str(v) for k, v in info.items()}
        # Fallback title from filename
        if not meta.title:
            meta.title = path.stem.replace("_", " ").replace("-", " ").title()
    except ImportError:
        logger.debug("PyPDF2 not installed, skipping PDF metadata")
    except Exception as e:
        logger.debug(f"PDF metadata error: {e}")


def _extract_docx_metadata(path: Path, meta: DocumentMetadata) -> None:
    """Extract metadata from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(str(path))
        core_props = doc.core_properties
        meta.author = core_props.author or ""
        meta.title = core_props.title or ""
        meta.created_date = str(core_props.created)[:10] if core_props.created else ""
        meta.page_count = len(doc.paragraphs)  # Approximate
        meta.word_count = sum(len(p.text.split()) for p in doc.paragraphs)
        if not meta.title:
            meta.title = path.stem.replace("_", " ").replace("-", " ").title()
    except ImportError:
        logger.debug("python-docx not installed, skipping DOCX metadata")
    except Exception as e:
        logger.debug(f"DOCX metadata error: {e}")


def _extract_text_metadata(path: Path, meta: DocumentMetadata) -> None:
    """Extract metadata from plain text files using regex."""
    try:
        content = path.read_text(encoding="utf-8", errors="replace")
        words = content.split()
        meta.word_count = len(words)

        # Extract title from first line / patterns
        for pattern in _TITLE_PATTERNS:
            match = pattern.search(content)
            if match:
                meta.title = match.group(1).strip()[:200]
                break

        if not meta.title:
            # Use first meaningful line
            for line in content.split("\n"):
                line = line.strip()
                if line and not line.startswith(("#", "//", "<!--", "/*")):
                    meta.title = line[:100]
                    break

        # Extract author
        for pattern in _AUTHOR_PATTERNS:
            match = pattern.search(content)
            if match:
                meta.author = match.group(1).strip()
                break

        # Extract date
        for pattern in _DATE_PATTERNS:
            match = pattern.search(content)
            if match:
                meta.created_date = match.group(1).strip()[:20]
                break

        # Extract keywords
        for pattern in _KEYWORD_PATTERNS:
            match = pattern.search(content)
            if match:
                kw_str = match.group(1).strip()
                meta.keywords = [k.strip() for k in kw_str.split(",") if k.strip()]

        # Detect language
        meta.language = _detect_language(content)

    except Exception as e:
        logger.debug(f"Text metadata error: {e}")


def _detect_language(text: str) -> str:
    """Simple language detection based on character analysis.

    Returns: "vi", "en", "mixed", or "unknown"
    """
    if not text or len(text) < 20:
        return "unknown"

    sample = text[:2000].lower()

    # Vietnamese-specific characters
    vn_chars = set("ăâđêôơưàảãáạằẳẵắặẩẫấậèẻẽéẹềểễếệìỉĩíịòỏõóọồổỗốộờởỡớợùủũúụừửữứựỳỷỹýỵ")
    vn_count = sum(1 for c in sample if c in vn_chars)

    if vn_count > 5:
        return "vi"
    return "en"


def _guess_mime_type(ext: str) -> str:
    """Guess MIME type from file extension."""
    mime_map = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".txt": "text/plain",
        ".md": "text/markdown",
        ".csv": "text/csv",
        ".json": "application/json",
        ".html": "text/html",
        ".htm": "text/html",
        ".xml": "text/xml",
        ".py": "text/x-python",
        ".js": "text/javascript",
        ".ts": "text/typescript",
    }
    return mime_map.get(ext.lower(), "application/octet-stream")


def _format_timestamp(ts: float) -> str:
    """Format a Unix timestamp to ISO date string."""
    try:
        return datetime.datetime.fromtimestamp(ts).strftime("%Y-%m-%d %H:%M:%S")
    except Exception:
        return ""
